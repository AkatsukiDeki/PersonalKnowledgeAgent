"""DOCX parser — extracts paragraphs and tables from Word documents."""

import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file, preserving heading hierarchy and tables."""
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            # It's a paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower()
            if "heading 1" in style:
                parts.append(f"# {text}")
            elif "heading 2" in style:
                parts.append(f"## {text}")
            elif "heading 3" in style:
                parts.append(f"### {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            # It's a table
            for tbl in doc.tables:
                if tbl._element is element:
                    rows_md = _table_to_markdown(tbl)
                    if rows_md:
                        parts.append(rows_md)
                    break

    result = "\n\n".join(parts)
    logger.info(f"[DOCXParser] Extracted {len(parts)} blocks, {len(result)} chars")
    return result


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table object to a Markdown table string."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
    body_lines = ["| " + " | ".join(r) + " |" for r in rows[1:]]

    return "\n".join([header, separator] + body_lines)
